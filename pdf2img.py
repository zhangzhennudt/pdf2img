# -*- coding: utf-8 -*-
"""
Created on Tue Sep 18 21:05:12 2018

@author: zhangzhennudt
@email:zhangzhennudt@126.com
"""

import fitz
import glob
import os
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from docx import Document

def rightinput(desc):
    flag=True
    while(flag):
        instr = input(desc)
        try:
            intnum = eval(instr)
            if type(intnum)==int:
                flag = False
        except:
            print('请输入正整数！')
            pass
    return intnum

def pdf2png(name):
    pdffile = glob.glob(name)[0]
    doc = fitz.open(pdffile)
    flag = rightinput("输入：1：全部页面；2：选择页面\t")
    if flag == 1:
        strat = 0
        totaling = doc.pageCount
    else:
        strat = rightinput('输入起始页面：') - 1
        totaling = rightinput('输入结束页面：')

    for pg in range(strat, totaling):
        page = doc[pg]
        zoom = int(100)
        rotate = int(0)
        trans = fitz.Matrix(zoom / 100.0, zoom / 100.0).preRotate(rotate)
        pm = page.getPixmap(matrix=trans, alpha=False)
        pm.writePNG('pdf2png/%s.png' % str(pg+1))

def sortfile(l):
    for i in range(len(l)):
        l[i] = l[i].split('.')
        l[i][0] = int(l[i][0][8:])
    l.sort()
    for i in range(len(l)):
        l[i][0] = str(l[i][0])
        l[i] = "pic2pdf\\"+l[i][0] + '.' + l[i][1]
    return l

def pdf2word(name):
    document = Document()
    # rb以二进制读模式打开本地pdf文件
    fn = open(name,'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
 
    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
 
    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource,laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource,device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out,"get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('test.txt','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style='ListBullet'    # 添加段落，样式为unordered list类型
                    )
                document.save('demo1.docx')  # 保存这个文档

def pic2pdf():
    doc = fitz.open()
    for img in sortfile(glob.glob("pic2pdf/*")):  # 读取图片，确保按文件名排序
        print(img)
        imgdoc = fitz.open(img)                 # 打开图片
        pdfbytes = imgdoc.convertToPDF()        # 使用图片创建单页的 PDF
        imgpdf = fitz.open("pdf", pdfbytes)
        doc.insertPDF(imgpdf)                   # 将当前页插入文档
    if os.path.exists("allimages.pdf"):
        os.remove("allimages.pdf")
    doc.save("allimages.pdf")                   # 保存pdf文件
    doc.close()

if __name__ == '__main__':
    name="盆底肌运动 - 副本.pptx.pdf"
#    pdf2png(name)
#    pic2pdf()
    pdf2word(name)